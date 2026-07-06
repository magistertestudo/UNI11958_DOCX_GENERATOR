from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel

from docx import Document

from pathlib import Path

import uuid
import os

app = FastAPI(
    title="UNI11958 DOCX Generator",
    version="1.0"
)


class RelazioneRequest(BaseModel):

    cliente: str = ""

    stabilimento: str = ""

    reparto: str = ""

    nome_ambiente: str = ""

    codice_ambiente: str = ""

    descrizione_ambiente: str = ""

    classificazione_ambiente: str = ""

    indice_globale: str = ""

    classe_globale: str = ""

    indice_globale_finale: str = ""

    classe_globale_finale: str = ""

    relazione_finale: str = ""

    conclusioni: str = ""


def replace_in_paragraph(paragraph, replacements):

    for key, value in replacements.items():

        if key in paragraph.text:

            paragraph.text = paragraph.text.replace(
                key,
                str(value)
            )


def replace_in_table(table, replacements):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                replace_in_paragraph(
                    paragraph,
                    replacements
                )


def replace_everywhere(doc, replacements):

    for paragraph in doc.paragraphs:

        replace_in_paragraph(
            paragraph,
            replacements
        )

    for table in doc.tables:

        replace_in_table(
            table,
            replacements
        )


@app.post("/generate-docx")
async def generate_docx(data: RelazioneRequest):

    template = Path("templates") / "MASTER_TEMPLATE_UNI11958_v1_1.docx"

    doc = Document(template)

    replacements = {

        "{{cliente}}": data.cliente,

        "{{stabilimento}}": data.stabilimento,

        "{{reparto}}": data.reparto,

        "{{nome_ambiente}}": data.nome_ambiente,

        "{{codice_ambiente}}": data.codice_ambiente,

        "{{descrizione_ambiente}}": data.descrizione_ambiente,

        "{{classificazione_ambiente}}": data.classificazione_ambiente,

        "{{indice_globale}}": data.indice_globale,

        "{{classe_globale}}": data.classe_globale,

        "{{indice_globale_finale}}": data.indice_globale_finale,

        "{{classe_globale_finale}}": data.classe_globale_finale,

        "{{relazione_finale}}": data.relazione_finale,

        "{{conclusioni}}": data.conclusioni

    }

    replace_everywhere(doc, replacements)

    output_folder = Path("outputs")

    output_folder.mkdir(exist_ok=True)

    filename = f"{uuid.uuid4().hex}.docx"

    output_file = output_folder / filename

    doc.save(output_file)

    return FileResponse(
        path=output_file,
        filename=f"Relazione_Tecnica_{data.nome_ambiente}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/")
def root():

    return {
        "service": "UNI11958 DOCX Generator",
        "status": "online",
        "version": "1.0"
    }